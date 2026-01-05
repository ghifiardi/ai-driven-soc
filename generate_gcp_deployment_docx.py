#!/usr/bin/env python3
"""
Generate GCP Threat Hunting Deployment Documentation in DOCX Format
Converts the markdown deployment documentation to a professionally formatted Word document.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import re


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def create_styled_table(doc, headers, rows, header_color="1F4E79", alt_row_color="D6E3F8"):
    """Create a professionally styled table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_shading(header_cells[i], header_color)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = str(cell_data)
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            if row_idx % 2 == 0:
                set_cell_shading(row.cells[i], alt_row_color)

    return table


def add_code_block(doc, code, language=""):
    """Add a code block with formatting."""
    para = doc.add_paragraph()
    para.style = 'No Spacing'
    run = para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 128)
    
    # Add light gray background
    para_format = para.paragraph_format
    para_format.left_indent = Inches(0.25)
    para_format.right_indent = Inches(0.25)
    para_format.space_before = Pt(6)
    para_format.space_after = Pt(6)
    
    return para


def parse_markdown_to_docx(md_file, docx_file):
    """Parse markdown file and create DOCX document."""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('GCP Threat Hunting Test Environment', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Deployment Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].bold = True
    
    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f'Version: 1.0 | Date: {datetime.now().strftime("%B %d, %Y")} | Environment: GCP Instance (xdgaisocapp01)')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.italic = True
    
    doc.add_paragraph()  # Spacing
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_rows = []
    table_headers = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip front matter
        if line.startswith('**Version:**') or line.startswith('**Date:**') or line.startswith('**Environment:**') or line.startswith('**Location:**'):
            i += 1
            continue
        
        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_block_lines)
                add_code_block(doc, code_text)
                code_block_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:], 1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], 4)
        
        # Horizontal rules
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            para = doc.add_paragraph(line[2:], style='List Bullet')
            para.runs[0].font.size = Pt(11)
        elif re.match(r'^\d+\. ', line):
            para = doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
            para.runs[0].font.size = Pt(11)
        
        # Bold text
        elif '**' in line:
            para = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                else:
                    para.add_run(part)
            para.runs[0].font.size = Pt(11)
        
        # Regular paragraphs
        elif line.strip() and not line.startswith('|'):
            para = doc.add_paragraph(line)
            para.runs[0].font.size = Pt(11)
        
        # Tables
        elif line.startswith('|') and '---' not in line:
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if not in_table:
                # First row is header
                table_headers = cells
                in_table = True
            else:
                table_rows.append(cells)
        
        # End table
        elif in_table and (not line.startswith('|') or '---' in line):
            if table_headers and table_rows:
                create_styled_table(doc, table_headers, table_rows)
            table_headers = []
            table_rows = []
            in_table = False
            if '---' not in line:
                continue
        
        i += 1
    
    # Handle any remaining table
    if in_table and table_headers and table_rows:
        create_styled_table(doc, table_headers, table_rows)
    
    # Add footer
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.add_run('GCP Threat Hunting Deployment Documentation | Version 1.0 | ' + datetime.now().strftime('%Y-%m-%d'))
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.italic = True
    
    # Save document
    doc.save(docx_file)
    print(f"✓ DOCX document created: {docx_file}")


if __name__ == "__main__":
    md_file = "docs/GCP_THREAT_HUNTING_DEPLOYMENT.md"
    docx_file = "docs/GCP_Threat_Hunting_Deployment_Documentation.docx"
    
    try:
        parse_markdown_to_docx(md_file, docx_file)
        print(f"\n✓ Successfully generated DOCX documentation")
        print(f"  Location: {docx_file}")
    except Exception as e:
        print(f"✗ Error generating DOCX: {e}")
        import traceback
        traceback.print_exc()

