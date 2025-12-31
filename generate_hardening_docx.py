"""
Generate Production Hardening Documentation in DOCX format.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_with_header(doc, headers, rows, col_widths=None):
    """Add a formatted table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "2E86AB")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            row.cells[col_idx].text = str(cell_text)

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table


def create_document():
    doc = Document()

    # Set document properties
    core_props = doc.core_properties
    core_props.author = "GATRA SOC Platform"
    core_props.title = "Production Hardening Guide"
    core_props.subject = "Security Operations Center Production Hardening"

    # Title
    title = doc.add_heading('GATRA SOC Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Production Hardening Guide', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Document info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(f"Version 1.0 | December 2024").italic = True

    doc.add_paragraph()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        "This document describes the production hardening implementation for the GATRA "
        "(Generalized Anomaly Detection & Triage Response Architecture) SOC Platform. "
        "The hardening effort transforms the development-ready codebase into a production-grade, "
        "multi-tenant security operations platform suitable for enterprise deployment."
    )

    doc.add_paragraph(
        "The initial production readiness assessment scored the system at 33/100. After implementing "
        "the hardening measures described in this document, the system achieves enterprise-grade "
        "standards for security, multi-tenancy, observability, and resilience."
    )

    # Key Improvements Table
    doc.add_heading('Key Improvements Summary', level=2)

    improvements = [
        ("Secret Management", "Hardcoded JWT secrets", "GCP Secret Manager + secure env vars"),
        ("Service Auth", "None", "HMAC-based service-to-service authentication"),
        ("Rate Limiting", "Configured but not enforced", "Token bucket with Redis backend"),
        ("Logging", "Basic text logs", "Structured JSON with correlation IDs"),
        ("Health Checks", "Simple /health endpoint", "Liveness + Readiness with dependency checks"),
        ("Resilience", "None", "Circuit breakers, retry, graceful shutdown"),
        ("Tenant Isolation", "Partial", "Full context propagation across services"),
    ]

    add_table_with_header(doc, ["Area", "Before", "After"], improvements, [1.5, 2, 2.5])
    doc.add_paragraph()

    # Architecture Section
    doc.add_heading('Architecture Overview', level=1)

    doc.add_heading('Production Components', level=2)

    components = [
        ("MSSP Platform", "mssp_platform_server_production.py", "Main orchestration service"),
        ("TAA Service", "taa_service_production.py", "Triage & Analysis Agent"),
        ("CRA Service", "cra_service_production.py", "Containment & Response Agent"),
        ("Production Module", "production/", "Shared hardening utilities"),
    ]

    add_table_with_header(doc, ["Component", "File", "Description"], components, [1.5, 2.5, 2])
    doc.add_paragraph()

    doc.add_heading('Production Module Structure', level=2)

    modules = [
        ("security.py", "Secret management, service authentication, input validation"),
        ("rate_limiter.py", "Token bucket rate limiting with Redis backend"),
        ("observability.py", "Structured logging, metrics, health checks"),
        ("resilience.py", "Circuit breakers, retry logic, graceful shutdown"),
        ("tenant_context.py", "Tenant context propagation across services"),
    ]

    add_table_with_header(doc, ["Module", "Purpose"], modules, [2, 4])
    doc.add_paragraph()

    # Security Hardening Section
    doc.add_heading('Security Hardening', level=1)

    doc.add_heading('1. Secret Management', level=2)
    doc.add_paragraph(
        "The SecretManager class provides a secure, hierarchical approach to secret retrieval "
        "with the following priority order:"
    )

    secret_priority = [
        ("1", "GCP Secret Manager", "Enterprise-grade secret storage (if GCP_SECRET_MANAGER_PROJECT is set)"),
        ("2", "Environment Variables", "For container deployments with orchestrator secrets"),
        ("3", "File-based secrets", "For development environments only"),
    ]

    add_table_with_header(doc, ["Priority", "Source", "Use Case"], secret_priority, [0.8, 1.7, 3.5])
    doc.add_paragraph()

    doc.add_heading('2. Service-to-Service Authentication', level=2)
    doc.add_paragraph(
        "All internal service calls are authenticated using HMAC-SHA256 signatures. This provides:"
    )

    bullets = doc.add_paragraph()
    bullets.add_run("• Message integrity verification (tampering prevention)\n")
    bullets.add_run("• Timestamp validation (replay attack prevention with 5-minute window)\n")
    bullets.add_run("• Service identification for audit trails\n")

    doc.add_heading('3. Input Validation', level=2)
    doc.add_paragraph(
        "The InputValidator class provides comprehensive input sanitization for all user inputs:"
    )

    validators = [
        ("validate_tenant_id()", "Alphanumeric + hyphens, 3-64 characters"),
        ("validate_alert_id()", "UUID format validation"),
        ("validate_ip_address()", "IPv4/IPv6 format validation"),
        ("sanitize_string()", "HTML/script injection prevention"),
    ]

    add_table_with_header(doc, ["Method", "Validation"], validators, [2, 4])
    doc.add_paragraph()

    doc.add_heading('4. API Key Security', level=2)
    doc.add_paragraph(
        "API keys are stored as salted SHA-256 hashes. Plain-text API keys are never stored "
        "in the database or logs."
    )

    # Tenant Isolation Section
    doc.add_heading('Tenant Isolation', level=1)

    doc.add_heading('1. Context Propagation', level=2)
    doc.add_paragraph(
        "Every request carries tenant context through the entire processing pipeline using "
        "Python's contextvars for async-safe propagation. The TenantContext object includes:"
    )

    context_fields = [
        ("tenant_id", "Unique tenant identifier"),
        ("tenant_config", "Tenant-specific configuration"),
        ("correlation_id", "Request tracking ID"),
        ("dataset_id", "Tenant's BigQuery dataset"),
    ]

    add_table_with_header(doc, ["Field", "Description"], context_fields, [2, 4])
    doc.add_paragraph()

    doc.add_heading('2. Middleware Integration', level=2)
    doc.add_paragraph(
        "FastAPI middleware automatically extracts tenant context from X-Tenant-ID headers "
        "and makes it available throughout the request lifecycle."
    )

    doc.add_heading('3. Scoped Clients', level=2)
    doc.add_paragraph(
        "All external service calls automatically include tenant context. The TenantAwareClient "
        "injects X-Tenant-ID headers, and TenantScopedBigQueryClient routes queries to "
        "tenant-specific datasets."
    )

    # Observability Section
    doc.add_heading('Observability', level=1)

    doc.add_heading('1. Structured Logging', level=2)
    doc.add_paragraph(
        "All logs are emitted as JSON for easy parsing by log aggregators (ELK, Splunk, etc.). "
        "Each log entry includes:"
    )

    log_fields = [
        ("timestamp", "ISO 8601 timestamp with timezone"),
        ("level", "Log level (DEBUG, INFO, WARNING, ERROR)"),
        ("service", "Service name"),
        ("correlation_id", "Request tracking ID"),
        ("tenant_id", "Tenant identifier"),
        ("message", "Log message"),
    ]

    add_table_with_header(doc, ["Field", "Description"], log_fields, [2, 4])
    doc.add_paragraph()

    doc.add_heading('2. Correlation IDs', level=2)
    doc.add_paragraph(
        "Every request is assigned a unique correlation ID that flows through all services. "
        "This enables end-to-end request tracing across the distributed system."
    )

    doc.add_heading('3. Metrics Collection', level=2)
    doc.add_paragraph("Prometheus-compatible metrics for monitoring:")

    metrics = [
        ("requests_total", "Counter", "Total requests by endpoint and status"),
        ("request_duration_seconds", "Histogram", "Request latency distribution"),
        ("rate_limit_exceeded", "Counter", "Rate limit violations by tenant"),
        ("circuit_breaker_state", "Gauge", "Circuit breaker status (0=closed, 1=open)"),
        ("active_requests", "Gauge", "Current in-flight requests"),
    ]

    add_table_with_header(doc, ["Metric", "Type", "Description"], metrics, [2.2, 1, 2.8])
    doc.add_paragraph()

    doc.add_heading('4. Health Checks', level=2)
    doc.add_paragraph("Kubernetes-compatible liveness and readiness probes:")

    health_endpoints = [
        ("/health/live", "GET", "Liveness probe - is the process alive?"),
        ("/health/ready", "GET", "Readiness probe - can we serve traffic?"),
        ("/health", "GET", "Basic health status"),
    ]

    add_table_with_header(doc, ["Endpoint", "Method", "Description"], health_endpoints, [1.5, 1, 3.5])
    doc.add_paragraph()

    # Resilience Patterns Section
    doc.add_heading('Resilience Patterns', level=1)

    doc.add_heading('1. Circuit Breaker', level=2)
    doc.add_paragraph(
        "Prevents cascade failures when downstream services are unhealthy. The circuit breaker "
        "has three states:"
    )

    states = [
        ("CLOSED", "Normal operation", "All requests allowed"),
        ("OPEN", "Failures exceeded threshold", "All requests rejected immediately"),
        ("HALF_OPEN", "Testing recovery", "Limited requests allowed to test service health"),
    ]

    add_table_with_header(doc, ["State", "Condition", "Behavior"], states, [1.2, 2, 2.8])
    doc.add_paragraph()

    doc.add_paragraph("Configuration parameters:")

    cb_config = [
        ("failure_threshold", "5", "Number of failures before opening circuit"),
        ("recovery_timeout", "30s", "Time before attempting recovery"),
        ("half_open_requests", "3", "Test requests allowed in half-open state"),
    ]

    add_table_with_header(doc, ["Parameter", "Default", "Description"], cb_config, [1.8, 1, 3.2])
    doc.add_paragraph()

    doc.add_heading('2. Retry Logic', level=2)
    doc.add_paragraph(
        "Automatic retry with exponential backoff for transient failures. Retries are only "
        "attempted for idempotent operations and specific exception types (ConnectionError, TimeoutError)."
    )

    retry_config = [
        ("max_retries", "3", "Maximum retry attempts"),
        ("base_delay", "1.0s", "Initial delay between retries"),
        ("max_delay", "30.0s", "Maximum delay cap"),
        ("exponential_base", "2", "Backoff multiplier"),
    ]

    add_table_with_header(doc, ["Parameter", "Default", "Description"], retry_config, [1.8, 1, 3.2])
    doc.add_paragraph()

    doc.add_heading('3. Graceful Shutdown', level=2)
    doc.add_paragraph(
        "Ensures in-flight requests complete before shutdown. When a SIGTERM signal is received:"
    )

    shutdown_steps = doc.add_paragraph()
    shutdown_steps.add_run("1. Stop accepting new requests\n")
    shutdown_steps.add_run("2. Wait for in-flight requests to complete (up to grace period)\n")
    shutdown_steps.add_run("3. Run registered cleanup tasks (close connections, flush buffers)\n")
    shutdown_steps.add_run("4. Exit process\n")

    doc.add_heading('4. Bulkhead Pattern', level=2)
    doc.add_paragraph(
        "Isolates resources to prevent one component from exhausting shared resources. "
        "For example, limiting concurrent BigQuery queries to 10 prevents query storms "
        "from blocking other operations."
    )

    # Deployment Guide Section
    doc.add_heading('Deployment Guide', level=1)

    doc.add_heading('Prerequisites', level=2)

    prereqs = doc.add_paragraph()
    prereqs.add_run("• Docker 20.10+\n")
    prereqs.add_run("• Docker Compose 2.0+\n")
    prereqs.add_run("• GCP Service Account with BigQuery Data Editor role\n")
    prereqs.add_run("• GCP Secret Manager access (optional but recommended)\n")
    prereqs.add_run("• Redis 7.0+ (included in docker-compose)\n")

    doc.add_heading('Quick Start', level=2)

    doc.add_paragraph("1. Create production environment file:")
    code1 = doc.add_paragraph()
    code1.add_run("cp .env.production.template .env.production").font.name = "Courier New"

    doc.add_paragraph("2. Generate secure secrets:")
    code2 = doc.add_paragraph()
    code2.add_run("export JWT_SECRET=$(openssl rand -hex 32)\n").font.name = "Courier New"
    code2.add_run("export SERVICE_API_KEY=$(openssl rand -hex 32)").font.name = "Courier New"

    doc.add_paragraph("3. Deploy with Docker Compose:")
    code3 = doc.add_paragraph()
    code3.add_run("docker-compose -f docker-compose.production.yml --env-file .env.production up -d").font.name = "Courier New"

    doc.add_paragraph("4. Verify deployment:")
    code4 = doc.add_paragraph()
    code4.add_run("curl http://localhost:8080/health/ready").font.name = "Courier New"

    doc.add_heading('Environment Variables', level=2)

    env_vars = [
        ("JWT_SECRET", "Yes", "JWT signing secret (min 32 chars)"),
        ("SERVICE_API_KEY", "Yes", "Service-to-service authentication key"),
        ("GCP_PROJECT_ID", "No", "GCP project for Secret Manager"),
        ("REDIS_URL", "Yes", "Redis connection URL"),
        ("LOG_LEVEL", "No", "Logging level (default: INFO)"),
        ("ENVIRONMENT", "No", "Environment name (default: production)"),
    ]

    add_table_with_header(doc, ["Variable", "Required", "Description"], env_vars, [2, 1, 3])
    doc.add_paragraph()

    # Operations Runbook Section
    doc.add_heading('Operations Runbook', level=1)

    doc.add_heading('Monitoring Alerts', level=2)

    alerts = [
        ("High Error Rate", ">5% 5xx responses", "Check service logs, circuit breaker state"),
        ("Rate Limit Exceeded", ">10 per minute", "Review tenant limits, check for abuse"),
        ("Circuit Breaker Open", "Any occurrence", "Check downstream service health"),
        ("High Latency", "p99 > 5s", "Scale services, optimize queries"),
    ]

    add_table_with_header(doc, ["Alert", "Threshold", "Action"], alerts, [1.8, 1.5, 2.7])
    doc.add_paragraph()

    doc.add_heading('Scaling Guidelines', level=2)

    scaling = [
        ("MSSP Platform", "70%", "80%", "Add replicas (horizontal scaling)"),
        ("TAA Service", "80%", "70%", "Add replicas + increase CPU limit"),
        ("CRA Service", "60%", "60%", "Add replicas"),
        ("Redis", "70%", "80%", "Upgrade instance size (vertical)"),
    ]

    add_table_with_header(doc, ["Service", "CPU Threshold", "Memory Threshold", "Action"], scaling, [1.5, 1.2, 1.4, 2])
    doc.add_paragraph()

    # Security Checklist Section
    doc.add_heading('Security Checklist', level=1)

    doc.add_heading('Pre-Deployment', level=2)

    pre_deploy = doc.add_paragraph()
    pre_deploy.add_run("☐ Generate unique JWT_SECRET (minimum 32 characters)\n")
    pre_deploy.add_run("☐ Generate unique SERVICE_API_KEY\n")
    pre_deploy.add_run("☐ Configure GCP Secret Manager (recommended)\n")
    pre_deploy.add_run("☐ Review and configure tenant rate limits\n")
    pre_deploy.add_run("☐ Enable TLS at load balancer\n")
    pre_deploy.add_run("☐ Configure WAF rules\n")
    pre_deploy.add_run("☐ Review service account permissions (least privilege)\n")

    doc.add_heading('Post-Deployment', level=2)

    post_deploy = doc.add_paragraph()
    post_deploy.add_run("☐ Verify health endpoints respond correctly\n")
    post_deploy.add_run("☐ Test rate limiting is enforced\n")
    post_deploy.add_run("☐ Verify tenant isolation (cross-tenant access blocked)\n")
    post_deploy.add_run("☐ Test circuit breaker behavior\n")
    post_deploy.add_run("☐ Verify structured logs are collected by log aggregator\n")
    post_deploy.add_run("☐ Test graceful shutdown behavior\n")
    post_deploy.add_run("☐ Run security scan on exposed endpoints\n")

    doc.add_heading('Ongoing', level=2)

    ongoing = doc.add_paragraph()
    ongoing.add_run("☐ Rotate secrets quarterly\n")
    ongoing.add_run("☐ Review access logs weekly\n")
    ongoing.add_run("☐ Monitor rate limit violations\n")
    ongoing.add_run("☐ Review circuit breaker events\n")
    ongoing.add_run("☐ Update dependencies monthly\n")

    # Glossary Section
    doc.add_heading('Glossary', level=1)

    glossary = [
        ("GATRA", "Generalized Anomaly Detection & Triage Response Architecture"),
        ("TAA", "Triage & Analysis Agent"),
        ("CRA", "Containment & Response Agent"),
        ("ADA", "Anomaly Detection Agent"),
        ("CLA", "Continuous Learning Agent"),
        ("MSSP", "Managed Security Service Provider"),
        ("EPS", "Events Per Second"),
        ("Circuit Breaker", "Pattern to prevent cascade failures by stopping requests to failing services"),
        ("Bulkhead", "Pattern to isolate resource consumption between components"),
        ("Correlation ID", "Unique identifier for tracking requests across distributed services"),
    ]

    add_table_with_header(doc, ["Term", "Definition"], glossary, [1.5, 4.5])
    doc.add_paragraph()

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("─" * 50)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"Document generated for GATRA SOC Platform v1.0").italic = True
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d')}").italic = True

    return doc


if __name__ == "__main__":
    doc = create_document()
    output_path = "docs/PRODUCTION_HARDENING_GUIDE.docx"
    doc.save(output_path)
    print(f"Documentation saved to: {output_path}")
