from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_integration_docx():
    doc = Document()
    
    # Setting up styles for a "Premium" look
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('AI-Driven SOC: Customer Integration Guide')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x2E, 0x50, 0x77) # Deep Blue

    doc.add_paragraph("\nWelcome to the AI-Driven SOC. This guide is designed to get your security data flowing into our AI engine with zero friction.\n")

    # Section 1: Easiest Path
    doc.add_heading('⚡ The Easiest Path: 30-Day Free Trial (Public API)', level=1)
    doc.add_paragraph("The fastest way to start is using our Secure Public API.")
    doc.add_paragraph("Time to integrated: < 5 minutes.", style='List Bullet')
    doc.add_paragraph("Complexity: Zero (Just code, no network changes).", style='List Bullet')
    doc.add_paragraph("Cost: Free (No billing required).", style='List Bullet')

    doc.add_heading('Step 1: Get your API Key', level=2)
    doc.add_paragraph("Your SOC administrator will provide you with a unique API Key and a Tenant ID.")

    doc.add_heading('Step 2: Send your first Event', level=2)
    doc.add_paragraph("Copy and paste this command into your terminal to verify connectivity.")
    
    # Code block simulation
    code_table = doc.add_table(rows=1, cols=1)
    code_table.style = 'Table Grid'
    cell = code_table.rows[0].cells[0]
    cell.paragraphs[0].text = '# Exchange API Key for a Token\nTOKEN=$(curl -s -X POST https://api.your-soc.com/api/v1/auth/token \\\n     -H "X-API-Key: YOUR_API_KEY" | jq -r .access_token)\n\n# Send a test event\ncurl -X POST https://api.your-soc.com/api/v1/events \\\n     -H "Authorization: Bearer $TOKEN" \\\n     -H "Content-Type: application/json" \\\n     -d \'{\n       "tenant_id": "YOUR_TENANT_ID",\n       "events": [{\n         "event_id": "test-001",\n         "type": "heartbeat",\n         "source": "onboarding_script"\n       }]\n     }\''
    for paragraph in cell.paragraphs:
        paragraph.style.font.name = 'Courier New'
        paragraph.style.font.size = Pt(9)

    doc.add_page_break()

    # Section 2: Engagement & Billing
    doc.add_heading('🗓️ Engagement & Billing Options', level=1)
    doc.add_paragraph("We offer two distinct engagement tiers tailored to your needs.")

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = '30-Day Trial'
    hdr_cells[2].text = '90-Day Engagement'

    data = [
        ('Onboarding Time', 'Minutes', '1-3 Days'),
        ('Connectivity', 'Public API (HTTPS)', 'Site-to-Site VPN (Private)'),
        ('Billing to You', 'Free ($0.00)', 'Standard Service Fee'),
        ('AI Benefits', 'Real-time monitoring', 'Continuous Learning (CLA)'),
        ('Best For', 'Testing & Demos', 'Enterprise Security')
    ]

    for i, (f, t, e) in enumerate(data):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = f
        row_cells[1].text = t
        row_cells[2].text = e

    doc.add_heading('💳 Is there a bill?', level=2)
    doc.add_paragraph("30-Day Trial: No. We use Sandbox mode. No cloud usage fees are incurred.", style='List Bullet')
    doc.add_paragraph("90-Day Engagement: Yes. Standard performance resources. Estimated ~$45/month in cloud infrastructure.", style='List Bullet')

    doc.add_page_break()

    # Section 3: Option B
    doc.add_heading('🛡️ Option B: Enterprise VPN (Private)', level=1)
    doc.add_paragraph("For 90-day engagements or highly regulated industries, we recommend a private connection.")
    doc.add_paragraph("Network: Site-to-Site IPSec VPN.", style='List Bullet')
    doc.add_paragraph("Auth: Mutual TLS (mTLS) + API Keys.", style='List Bullet')
    doc.add_paragraph("Setup: Our engineers will coordinate with your IT team.", style='List Bullet')

    doc.add_heading('🧪 Data Format (JSON)', level=1)
    doc.add_paragraph("Our engine accepts standard JSON. You don't need to change your logs; just wrap them in our template.")

    doc.add_page_break()

    # Section 4: Python Sample
    doc.add_heading('🐍 Appendix: Python Ingestion Script', level=1)
    doc.add_paragraph("For customers using Python, we provide this ready-to-use script for automated ingestion.")
    
    python_code = (
        "import requests\n"
        "import json\n"
        "import time\n\n"
        "# --- CONFIGURATION ---\n"
        "API_KEY = \"YOUR_API_KEY\"\n"
        "TENANT_ID = \"YOUR_TENANT_ID\"\n"
        "BASE_URL = \"https://api.your-soc.com/api/v1\"\n\n"
        "def get_token():\n"
        "    response = requests.post(f\"{BASE_URL}/auth/token\", headers={\"X-API-Key\": API_KEY})\n"
        "    return response.json()[\"access_token\"]\n\n"
        "def send_events(token, events):\n"
        "    requests.post(\n"
        "        f\"{BASE_URL}/events\",\n"
        "        headers={\"Authorization\": f\"Bearer {token}\", \"Content-Type\": \"application/json\"},\n"
        "        json={\"tenant_id\": TENANT_ID, \"events\": events}\n"
        "    )\n"
    )
    
    code_table_py = doc.add_table(rows=1, cols=1)
    code_table_py.style = 'Table Grid'
    cell_py = code_table_py.rows[0].cells[0]
    cell_py.paragraphs[0].text = python_code
    for paragraph in cell_py.paragraphs:
        paragraph.style.font.name = 'Courier New'
        paragraph.style.font.size = Pt(8)

    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer.paragraphs[0].text = "AI-Driven SOC Platform | Confidential"
    footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    filename = 'AI_SOC_Customer_Integration_Guide.docx'
    doc.save(filename)
    print(f"✅ Document saved as {filename}")
    return os.path.abspath(filename)

if __name__ == "__main__":
    create_integration_docx()
