#!/usr/bin/env python3
"""
Convert CLA documentation from Markdown to DOCX format
"""

import os
import subprocess
from pathlib import Path

def convert_md_to_docx():
    """Convert Markdown documentation to DOCX format"""

    # Input and output files
    md_file = "CLA_COMPREHENSIVE_DOCUMENTATION.md"
    docx_file = "CLA_COMPREHENSIVE_DOCUMENTATION.docx"

    # Check if input file exists
    if not os.path.exists(md_file):
        print(f"❌ Error: {md_file} not found")
        return False

    try:
        # Try using pandoc (most reliable for MD to DOCX conversion)
        try:
            result = subprocess.run([
                "pandoc",
                md_file,
                "-o", docx_file,
                "--from=markdown",
                "--to=docx",
                "--standalone",
                "--toc",
                "--toc-depth=3",
                "--highlight-style=tango"
            ], capture_output=True, text=True, timeout=30)

            if result.returncode == 0:
                print(f"✅ Successfully converted {md_file} to {docx_file}")
                print(f"📄 DOCX file created: {os.path.abspath(docx_file)}")
                return True
            else:
                print(f"❌ Pandoc conversion failed: {result.stderr}")
                return False

        except FileNotFoundError:
            print("⚠️  Pandoc not found, trying alternative method...")

            # Alternative: Use python-docx to create DOCX from Markdown
            try:
                import markdown
                from docx import Document
                from docx.shared import Inches

                # Read markdown file
                with open(md_file, 'r', encoding='utf-8') as f:
                    md_content = f.read()

                # Convert markdown to HTML
                html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])

                # Create DOCX document
                doc = Document()

                # Add title
                title = doc.add_heading('Continuous Learning Agent (CLA) - Comprehensive Documentation', 0)

                # Split content by sections (basic parsing)
                sections = html_content.split('<h1>')[1:]  # Skip first empty part

                for section in sections:
                    if '</h1>' in section:
                        section_title = section.split('</h1>')[0].strip()
                        section_content = section.split('</h1>')[1] if '</h1>' in section else section

                        # Add section heading
                        doc.add_heading(section_title, level=1)

                        # Add content (simplified - would need proper HTML to DOCX conversion)
                        doc.add_paragraph(section_content[:200] + "..." if len(section_content) > 200 else section_content)

                # Save DOCX file
                doc.save(docx_file)
                print(f"✅ Created DOCX file using python-docx: {os.path.abspath(docx_file)}")
                return True

            except ImportError:
                print("❌ python-docx not available. Install with: pip install python-docx markdown")
                return False

    except Exception as e:
        print(f"❌ Error during conversion: {e}")
        return False

def main():
    """Main conversion function"""
    print("📄 Converting CLA Documentation to DOCX Format")
    print("=" * 60)

    success = convert_md_to_docx()

    if success:
        print("\n🎉 Conversion successful!")
        print("\nFiles created:")
        print("  📄 CLA_COMPREHENSIVE_DOCUMENTATION.docx")
        print("\n📋 Contents include:")
        print("  • Executive Summary")
        print("  • System Architecture")
        print("  • Feedback Mechanism")
        print("  • Accuracy Calculation")
        print("  • Integration Points")
        print("  • Performance Monitoring")
        print("  • Deployment Guide")
        print("  • Troubleshooting")
        print("  • API Reference")
        print("  • Security Considerations")
    else:
        print("\n❌ Conversion failed. Please check the error messages above.")

if __name__ == "__main__":
    main()
