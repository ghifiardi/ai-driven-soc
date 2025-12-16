#!/usr/bin/env python3
"""
Documentation Generator for Hybrid Cyber Defense Agent
Converts Markdown documentation to DOCX format for production deployment
"""

import os
import re
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("Installing required dependencies...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add color and underline
    color = OxmlElement('w:color')
    color.set(qn('w:val'), "0563C1")
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def parse_markdown_to_docx(markdown_file, output_file):
    """Convert Markdown file to DOCX format"""
    
    # Read markdown file
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create new document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Hybrid Cyber Defense Agent - Production Deployment Guide"
    doc.core_properties.author = "AI-Driven SOC Development Team"
    doc.core_properties.subject = "Cybersecurity AI Agent Deployment Documentation"
    doc.core_properties.comments = "Comprehensive deployment guide for Hybrid Cyber Defense Agent system"
    
    # Add title page
    title = doc.add_heading('Hybrid Cyber Defense Agent', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Production Deployment Guide', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add document info
    doc.add_paragraph(f'Version: 1.0')
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %Y")}')
    doc.add_paragraph(f'Author: AI-Driven SOC Development Team')
    doc.add_paragraph(f'Classification: Technical Documentation')
    
    # Add page break
    doc.add_page_break()
    
    # Add table of contents
    toc_heading = doc.add_heading('Table of Contents', level=1)
    
    # Parse content line by line
    lines = content.split('\n')
    in_code_block = False
    code_language = ""
    current_heading_level = 0
    
    for line in lines:
        line = line.strip()
        
        # Skip empty lines
        if not line:
            doc.add_paragraph()
            continue
            
        # Handle code blocks
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_language = line[3:].strip()
                continue
            else:
                in_code_block = False
                continue
        
        if in_code_block:
            # Add code paragraph
            code_para = doc.add_paragraph(line)
            code_para.style = 'Code'
            continue
        
        # Handle headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            heading_text = line.lstrip('# ').strip()
            
            # Create heading
            heading = doc.add_heading(heading_text, level=level)
            
            # Add to TOC if it's a main section
            if level <= 2:
                toc_para = doc.add_paragraph(f'{"  " * (level - 1)}• {heading_text}')
                toc_para.style = 'List Bullet'
            
            current_heading_level = level
            continue
        
        # Handle lists
        if line.startswith('- ') or line.startswith('* '):
            list_text = line[2:].strip()
            para = doc.add_paragraph(list_text, style='List Bullet')
            continue
        
        if line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            list_text = re.sub(r'^\d+\. ', '', line)
            para = doc.add_paragraph(list_text, style='List Number')
            continue
        
        # Handle horizontal rules
        if line.startswith('---'):
            doc.add_paragraph('_' * 50)
            continue
        
        # Handle links
        if '[' in line and '](' in line:
            # Simple link handling
            para = doc.add_paragraph()
            parts = line.split('[')
            if parts[0]:
                para.add_run(parts[0])
            
            for part in parts[1:]:
                if '](' in part and ')' in part:
                    link_text, link_url = part.split('](')
                    link_url = link_url.split(')')[0]
                    para.add_run(link_text)
                    add_hyperlink(para, link_text, link_url)
                    remaining = part.split(')', 1)[1] if ')' in part else ''
                    if remaining:
                        para.add_run(remaining)
                else:
                    para.add_run('[' + part)
            continue
        
        # Handle bold text
        if '**' in line:
            para = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    para.add_run(part)
                else:
                    run = para.add_run(part)
                    run.bold = True
            continue
        
        # Handle italic text
        if '*' in line and not line.startswith('*'):
            para = doc.add_paragraph()
            parts = line.split('*')
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    para.add_run(part)
                else:
                    run = para.add_run(part)
                    run.italic = True
            continue
        
        # Handle inline code
        if '`' in line:
            para = doc.add_paragraph()
            parts = line.split('`')
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    para.add_run(part)
                else:
                    run = para.add_run(part)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            continue
        
        # Regular paragraph
        if line:
            para = doc.add_paragraph(line)
    
    # Add footer with page numbers
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Hybrid Cyber Defense Agent Deployment Guide - Page "
    
    # Save document
    doc.save(output_file)
    print(f"✅ DOCX document created: {output_file}")

def create_deployment_package():
    """Create a complete deployment package with all documentation"""
    
    print("📦 Creating Hybrid Cyber Defense Agent Deployment Package...")
    
    # Create deployment directory
    deployment_dir = Path("deployment_package")
    deployment_dir.mkdir(exist_ok=True)
    
    # Copy documentation files
    docs_dir = deployment_dir / "documentation"
    docs_dir.mkdir(exist_ok=True)
    
    # List of documentation files to include
    doc_files = [
        "HYBRID_DEFENSE_AGENT_DEPLOYMENT_GUIDE.md",
        "HYBRID_DEFENSE_AGENT_TECHNICAL_SPEC.md",
        "HYBRID_DEFENSE_DASHBOARD_GUIDE.md",
        "DASHBOARD_STATUS.md"
    ]
    
    # Copy markdown files
    for doc_file in doc_files:
        if Path(doc_file).exists():
            import shutil
            shutil.copy2(doc_file, docs_dir)
            print(f"📄 Copied: {doc_file}")
    
    # Generate DOCX versions
    for doc_file in doc_files:
        if Path(doc_file).exists():
            docx_file = docs_dir / f"{Path(doc_file).stem}.docx"
            try:
                parse_markdown_to_docx(doc_file, str(docx_file))
                print(f"📝 Generated DOCX: {docx_file}")
            except Exception as e:
                print(f"❌ Error generating DOCX for {doc_file}: {e}")
    
    # Copy configuration files
    config_dir = deployment_dir / "config"
    config_dir.mkdir(exist_ok=True)
    
    config_files = [
        "config/hybrid_defense_config.json",
        "config/hybrid_agent_card.json"
    ]
    
    for config_file in config_files:
        if Path(config_file).exists():
            import shutil
            shutil.copy2(config_file, config_dir)
            print(f"⚙️ Copied config: {config_file}")
    
    # Copy service files
    service_dir = deployment_dir / "services"
    service_dir.mkdir(exist_ok=True)
    
    service_files = [
        "hybrid-defense.service",
        "hybrid-defense-dashboard.service"
    ]
    
    for service_file in service_files:
        if Path(service_file).exists():
            import shutil
            shutil.copy2(service_file, service_dir)
            print(f"🔧 Copied service: {service_file}")
    
    # Copy deployment scripts
    scripts_dir = deployment_dir / "scripts"
    scripts_dir.mkdir(exist_ok=True)
    
    script_files = [
        "deploy_hybrid_dashboard.sh",
        "test_dashboard_connection.py",
        "simple_validation.py"
    ]
    
    for script_file in script_files:
        if Path(script_file).exists():
            import shutil
            shutil.copy2(script_file, scripts_dir)
            print(f"🚀 Copied script: {script_file}")
    
    # Create deployment checklist
    checklist_content = """# Hybrid Cyber Defense Agent - Deployment Checklist

## Pre-Deployment
- [ ] Review system requirements
- [ ] Set up Google Cloud project
- [ ] Configure service accounts
- [ ] Install dependencies
- [ ] Review configuration files

## Deployment
- [ ] Deploy agent service
- [ ] Deploy dashboard service
- [ ] Configure Pub/Sub topics
- [ ] Test agent connectivity
- [ ] Verify dashboard access

## Post-Deployment
- [ ] Monitor system health
- [ ] Test alert processing
- [ ] Verify circuit breakers
- [ ] Check log aggregation
- [ ] Validate metrics collection

## Validation
- [ ] Health check endpoints
- [ ] A2A protocol compliance
- [ ] Pub/Sub message flow
- [ ] Dashboard functionality
- [ ] Performance benchmarks

## Documentation
- [ ] Update runbooks
- [ ] Train operations team
- [ ] Document procedures
- [ ] Create monitoring dashboards
- [ ] Establish escalation procedures
"""
    
    with open(deployment_dir / "DEPLOYMENT_CHECKLIST.md", 'w') as f:
        f.write(checklist_content)
    
    print(f"✅ Created deployment checklist")
    
    # Create README for deployment package
    readme_content = f"""# Hybrid Cyber Defense Agent - Deployment Package

This package contains all necessary files and documentation for deploying the Hybrid Cyber Defense Agent to a production environment.

## Package Contents

### Documentation
- `HYBRID_DEFENSE_AGENT_DEPLOYMENT_GUIDE.md` - Comprehensive deployment guide
- `HYBRID_DEFENSE_AGENT_DEPLOYMENT_GUIDE.docx` - DOCX version of deployment guide
- `HYBRID_DEFENSE_AGENT_TECHNICAL_SPEC.md` - Technical specifications
- `HYBRID_DEFENSE_AGENT_TECHNICAL_SPEC.docx` - DOCX version of technical spec
- `DEPLOYMENT_CHECKLIST.md` - Step-by-step deployment checklist

### Configuration
- `hybrid_defense_config.json` - Agent configuration
- `hybrid_agent_card.json` - A2A agent card

### Services
- `hybrid-defense.service` - Systemd service file for agent
- `hybrid-defense-dashboard.service` - Systemd service file for dashboard

### Scripts
- `deploy_hybrid_dashboard.sh` - Dashboard deployment script
- `test_dashboard_connection.py` - Connection testing script
- `simple_validation.py` - System validation script

## Quick Start

1. Review the deployment guide: `documentation/HYBRID_DEFENSE_AGENT_DEPLOYMENT_GUIDE.md`
2. Follow the deployment checklist: `DEPLOYMENT_CHECKLIST.md`
3. Configure your environment using the provided configuration files
4. Deploy using the provided service files and scripts

## Support

For technical support and questions:
- Review the technical specification document
- Check the troubleshooting section in the deployment guide
- Contact the development team

## Package Information
- Created: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
- Version: 1.0
- Author: AI-Driven SOC Development Team
"""
    
    with open(deployment_dir / "README.md", 'w') as f:
        f.write(readme_content)
    
    print(f"📋 Created package README")
    
    # Create summary
    print(f"\n🎉 Deployment package created successfully!")
    print(f"📁 Location: {deployment_dir.absolute()}")
    print(f"📊 Contents:")
    print(f"   📄 Documentation: {len(list(docs_dir.glob('*')))} files")
    print(f"   ⚙️ Configuration: {len(list(config_dir.glob('*')))} files")
    print(f"   🔧 Services: {len(list(service_dir.glob('*')))} files")
    print(f"   🚀 Scripts: {len(list(scripts_dir.glob('*')))} files")
    
    return deployment_dir

def main():
    """Main function to generate documentation package"""
    
    print("🚀 Hybrid Cyber Defense Agent - Documentation Generator")
    print("=" * 60)
    
    # Check if we're in the right directory
    if not Path("hybrid_cyber_defense_agent.py").exists():
        print("❌ Error: Please run this script from the project root directory")
        return
    
    try:
        # Create deployment package
        deployment_dir = create_deployment_package()
        
        print(f"\n✅ Documentation package created successfully!")
        print(f"📦 Package location: {deployment_dir.absolute()}")
        print(f"\n📋 Next steps:")
        print(f"   1. Review the documentation in the 'documentation' folder")
        print(f"   2. Follow the DEPLOYMENT_CHECKLIST.md")
        print(f"   3. Configure your environment")
        print(f"   4. Deploy to production")
        
    except Exception as e:
        print(f"❌ Error creating documentation package: {e}")
        return

if __name__ == "__main__":
    main()


















