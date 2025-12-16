#!/usr/bin/env python3
"""
Simple DOCX Documentation Creator for Hybrid Cyber Defense Agent
Creates professional DOCX documents from markdown files
"""

import os
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Installing required dependencies...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

def create_deployment_guide_docx():
    """Create DOCX version of the deployment guide"""
    
    # Create new document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Hybrid Cyber Defense Agent - Production Deployment Guide"
    doc.core_properties.author = "AI-Driven SOC Development Team"
    doc.core_properties.subject = "Cybersecurity AI Agent Deployment Documentation"
    doc.core_properties.comments = "Comprehensive deployment guide for Hybrid Cyber Defense Agent system"
    
    # Add title page
    title = doc.add_heading('Hybrid Cyber Defense Agent', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Production Deployment Guide', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add document info
    doc.add_paragraph(f'Version: 1.0')
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %Y")}')
    doc.add_paragraph(f'Author: AI-Driven SOC Development Team')
    doc.add_paragraph(f'Classification: Technical Documentation')
    
    # Add page break
    doc.add_page_break()
    
    # Add table of contents
    doc.add_heading('Table of Contents', level=1)
    
    toc_items = [
        'System Overview',
        'Architecture',
        'Prerequisites',
        'Installation & Setup',
        'Configuration',
        'Deployment',
        'Monitoring & Operations',
        'Troubleshooting',
        'Security Considerations',
        'Performance Optimization',
        'Maintenance',
        'Appendix'
    ]
    
    for item in toc_items:
        para = doc.add_paragraph(f'• {item}')
        para.style = 'List Bullet'
    
    # Add page break
    doc.add_page_break()
    
    # System Overview Section
    doc.add_heading('System Overview', level=1)
    
    overview_text = """
The Hybrid Cyber Defense Agent is an AI-powered cybersecurity system that combines:
• DQN (Deep Q-Network) for fast decision-making
• Google Gemini for explainable AI and natural language processing
• Pub/Sub integration for real-time alert processing
• A2A Protocol compliance for agent interoperability

Key Features:
• Real-time threat analysis and response
• Explainable AI with natural language explanations
• Circuit breaker patterns for resilience
• Comprehensive monitoring and alerting
• RESTful API for integration
• Streamlit dashboard for visualization
"""
    
    for line in overview_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())
    
    # Architecture Section
    doc.add_heading('Architecture', level=1)
    
    arch_text = """
The system follows a modular architecture with the following components:

Hybrid Defense Agent:
• DQN Engine: Fast threat classification and action recommendation
• Gemini Integration: Natural language explanations and reasoning
• Pub/Sub Handler: Real-time alert ingestion and response publishing
• Circuit Breakers: Resilience and fault tolerance
• A2A Compliance: Agent-to-agent communication protocol

Dashboard:
• Real-time Monitoring: Agent status, metrics, and performance
• Alert Visualization: Live threat processing and response tracking
• System Health: Component status and circuit breaker states
• Analytics: DQN performance and Gemini usage statistics
"""
    
    for line in arch_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())
    
    # Prerequisites Section
    doc.add_heading('Prerequisites', level=1)
    
    prereq_text = """
System Requirements:
• OS: Linux (Ubuntu 20.04+), macOS, or Windows 10+
• Python: 3.8+ (recommended 3.9+)
• Memory: 4GB+ RAM
• Storage: 10GB+ free space
• Network: Internet connectivity for Google Cloud services

Google Cloud Requirements:
• Project: Valid GCP project with billing enabled
• APIs: Pub/Sub, Vertex AI, Cloud Storage
• Service Account: With appropriate permissions
• Authentication: Application Default Credentials or service account key
"""
    
    for line in prereq_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())
    
    # Installation Section
    doc.add_heading('Installation & Setup', level=1)
    
    install_text = """
1. Environment Setup:
   • Create virtual environment: python3 -m venv hybrid-defense-env
   • Activate environment: source hybrid-defense-env/bin/activate
   • Upgrade pip: pip install --upgrade pip

2. Install Dependencies:
   • Install from requirements.txt: pip install -r requirements.txt
   • Or install individually: pip install torch torchvision vertexai google-cloud-pubsub fastapi uvicorn streamlit gymnasium

3. Google Cloud Setup:
   • Install Google Cloud CLI
   • Authenticate: gcloud auth login
   • Set project: gcloud config set project YOUR_PROJECT_ID
   • Enable APIs: gcloud services enable pubsub.googleapis.com aiplatform.googleapis.com
"""
    
    for line in install_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            elif line.strip().startswith(('1.', '2.', '3.')):
                doc.add_paragraph(line.strip(), style='List Number')
            else:
                doc.add_paragraph(line.strip())
    
    # Configuration Section
    doc.add_heading('Configuration', level=1)
    
    config_text = """
The system uses a JSON configuration file located at config/hybrid_defense_config.json.

Key configuration sections:
• agent: Agent-specific settings (ID, port, logging)
• dqn_model: DQN model configuration (state size, action size, model path)
• gemini: Gemini AI configuration (project ID, model name, parameters)
• pubsub: Pub/Sub configuration (topics, subscriptions, settings)
• circuit_breakers: Circuit breaker thresholds and timeouts

Environment Variables:
• GOOGLE_APPLICATION_CREDENTIALS: Path to service account key
• GOOGLE_CLOUD_PROJECT: Google Cloud project ID
• HYBRID_DEFENSE_CONFIG_PATH: Path to configuration file
• HYBRID_DEFENSE_LOG_LEVEL: Logging level (DEBUG, INFO, WARNING, ERROR)
"""
    
    for line in config_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())
    
    # Deployment Section
    doc.add_heading('Deployment', level=1)
    
    deploy_text = """
1. Development Deployment:
   • Start agent: python3 hybrid_cyber_defense_agent.py
   • Start dashboard: python3 -m streamlit run hybrid_defense_dashboard.py --server.port=8529

2. Production Deployment with systemd:
   • Create service file: /etc/systemd/system/hybrid-defense-agent.service
   • Enable service: sudo systemctl enable hybrid-defense-agent
   • Start service: sudo systemctl start hybrid-defense-agent
   • Check status: sudo systemctl status hybrid-defense-agent

3. Docker Deployment:
   • Use provided Dockerfile
   • Build image: docker build -t hybrid-defense-agent .
   • Run container: docker run -p 8083:8083 hybrid-defense-agent

4. Kubernetes Deployment:
   • Use provided Kubernetes manifests
   • Deploy: kubectl apply -f k8s/
   • Check pods: kubectl get pods
"""
    
    for line in deploy_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            elif line.strip().startswith(('1.', '2.', '3.', '4.')):
                doc.add_paragraph(line.strip(), style='List Number')
            else:
                doc.add_paragraph(line.strip())
    
    # Monitoring Section
    doc.add_heading('Monitoring & Operations', level=1)
    
    monitor_text = """
Health Checks:
• Agent health: curl http://localhost:8083/health
• Agent status: curl http://localhost:8083/status
• Dashboard health: curl http://localhost:8529/healthz

Logging:
• View agent logs: sudo journalctl -u hybrid-defense-agent -f
• View dashboard logs: sudo journalctl -u hybrid-defense-dashboard -f
• View specific log levels: sudo journalctl -u hybrid-defense-agent --since "1 hour ago" | grep ERROR

Metrics Collection:
The agent provides comprehensive metrics via the /status endpoint:
• Performance Metrics: DQN inference duration, Gemini API response times
• Business Metrics: Alerts processed, actions recommended, explanations generated
• Error Metrics: Error counts by type, circuit breaker states
• System Metrics: Uptime, memory usage, connection status
"""
    
    for line in monitor_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())
    
    # Troubleshooting Section
    doc.add_heading('Troubleshooting', level=1)
    
    trouble_text = """
Common Issues and Solutions:

1. Google Cloud Authentication:
   Problem: 404 Requested project not found or user does not have access
   Solution: Verify project ID, re-authenticate, check service account permissions

2. Pub/Sub Connection Issues:
   Problem: Connection refused or Timeout
   Solution: Check topic existence, check subscription, test connectivity

3. DQN Model Loading:
   Problem: No pre-trained model found
   Solution: Check model file exists, train model if needed, verify model loading

4. Gemini API Issues:
   Problem: Vertex AI API not available
   Solution: Check API enablement, verify quota, test API access

Debug Mode:
Enable debug logging: export HYBRID_DEFENSE_LOG_LEVEL=DEBUG
"""
    
    for line in trouble_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            elif line.strip().startswith(('1.', '2.', '3.', '4.')):
                doc.add_paragraph(line.strip(), style='List Number')
            else:
                doc.add_paragraph(line.strip())
    
    # Security Section
    doc.add_heading('Security Considerations', level=1)
    
    security_text = """
Authentication & Authorization:
• Use service accounts with minimal required permissions
• Rotate service account keys regularly
• Implement API key authentication for external access
• Use HTTPS in production

Network Security:
• Deploy behind load balancer with SSL termination
• Use VPC for internal communication
• Implement firewall rules to restrict access
• Monitor network traffic for anomalies

Data Protection:
• Encrypt sensitive data at rest
• Use secure communication channels
• Implement data retention policies
• Regular security audits
"""
    
    for line in security_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())
    
    # Performance Section
    doc.add_heading('Performance Optimization', level=1)
    
    perf_text = """
DQN Model Optimization:
• Use GPU acceleration when available
• Implement model quantization
• Batch processing for multiple alerts
• Model caching and preloading

Gemini API Optimization:
• Implement request batching
• Use connection pooling
• Cache frequent explanations
• Implement rate limiting

Pub/Sub Optimization:
• Increase message batch size
• Optimize acknowledgment timing
• Use flow control settings
• Monitor message backlog

System Optimization:
• Use SSD storage for models
• Increase memory allocation
• Optimize Python GIL usage
• Implement connection pooling
"""
    
    for line in perf_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())
    
    # Maintenance Section
    doc.add_heading('Maintenance', level=1)
    
    maintenance_text = """
Regular Tasks:

Daily:
• Monitor system health and performance
• Check error logs and circuit breaker states
• Verify Pub/Sub message flow
• Review alert processing metrics

Weekly:
• Update DQN model with new training data
• Review and optimize Gemini prompts
• Analyze performance trends
• Update documentation

Monthly:
• Security audit and vulnerability assessment
• Performance tuning and optimization
• Backup and disaster recovery testing
• Capacity planning review

Model Updates:
• Retrain DQN model: python3 train_dqn_model.py --episodes=1000
• Validate new model: python3 simple_validation.py
• Deploy new model: sudo systemctl restart hybrid-defense-agent
"""
    
    for line in maintenance_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())
    
    # Appendix Section
    doc.add_heading('Appendix', level=1)
    
    appendix_text = """
API Endpoints:
• GET /health - Health check
• GET /status - Comprehensive status and metrics
• POST /a2a/process_alert - Process alert via A2A protocol
• GET /agent-card - A2A agent card

Error Codes:
• 200: Success
• 400: Bad Request
• 401: Unauthorized
• 403: Forbidden
• 404: Not Found
• 500: Internal Server Error
• 503: Service Unavailable

Support Contacts:
• Technical Support: [Your support email]
• Documentation: [Your documentation URL]
• Issue Tracking: [Your issue tracking system]
• Emergency Contact: [Your emergency contact]

Document Version: 1.0
Last Updated: {datetime.now().strftime("%B %Y")}
Author: Hybrid Defense Team
Review Date: [Review Date]
"""
    
    for line in appendix_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())
    
    # Save document
    output_file = "deployment_package/documentation/HYBRID_DEFENSE_AGENT_DEPLOYMENT_GUIDE.docx"
    doc.save(output_file)
    print(f"✅ DOCX document created: {output_file}")

def create_technical_spec_docx():
    """Create DOCX version of the technical specification"""
    
    # Create new document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Hybrid Cyber Defense Agent - Technical Specification"
    doc.core_properties.author = "AI-Driven SOC Development Team"
    doc.core_properties.subject = "Technical Specification for Hybrid Cyber Defense Agent"
    doc.core_properties.comments = "Comprehensive technical specification for Hybrid Cyber Defense Agent system"
    
    # Add title page
    title = doc.add_heading('Hybrid Cyber Defense Agent', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Technical Specification', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add document info
    doc.add_paragraph(f'Version: 1.0')
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %Y")}')
    doc.add_paragraph(f'Author: AI-Driven SOC Development Team')
    doc.add_paragraph(f'Classification: Technical Documentation')
    
    # Add page break
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    
    summary_text = """
The Hybrid Cyber Defense Agent is a sophisticated AI-powered cybersecurity system that integrates Deep Q-Network (DQN) machine learning with Google Gemini's natural language processing capabilities. This system provides real-time threat analysis, automated response recommendations, and explainable AI for cybersecurity operations.

Key Technical Features:
• DQN-based threat classification with <100ms inference time
• Google Gemini integration for explainable AI
• Pub/Sub-based real-time alert processing
• Circuit breaker patterns for system resilience
• A2A protocol compliance for agent interoperability
• Comprehensive monitoring and alerting
• RESTful API for external integration
"""
    
    for line in summary_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())
    
    # System Architecture
    doc.add_heading('System Architecture', level=1)
    
    arch_text = """
The system follows a modular, microservices-based architecture:

Core Components:
• Hybrid Defense Agent: Main processing engine
• DQN Engine: Machine learning inference
• Gemini Integration: Natural language processing
• Pub/Sub Handler: Message processing
• Circuit Breakers: Fault tolerance
• Dashboard: Monitoring and visualization

Architecture Benefits:
• Scalability: Horizontal scaling capability
• Reliability: Circuit breaker patterns
• Maintainability: Modular design
• Extensibility: Plugin architecture
• Observability: Comprehensive monitoring
"""
    
    for line in arch_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())
    
    # Technical Specifications
    doc.add_heading('Technical Specifications', level=1)
    
    tech_spec_text = """
System Requirements:
• OS: Linux (Ubuntu 20.04+), macOS, or Windows 10+
• Python: 3.8+ (recommended 3.9+)
• Memory: 4GB+ RAM (8GB+ recommended)
• Storage: 10GB+ free space (50GB+ recommended)
• Network: Internet connectivity for Google Cloud services

Performance Requirements:
• DQN Inference: <100ms (95th percentile)
• Gemini API Call: <10 seconds
• Total Alert Processing: <15 seconds (95th percentile)
• Health Check: <1 second
• Status Endpoint: <2 seconds

Throughput Requirements:
• Alerts per Second: 10 (sustained), 50 (burst)
• Concurrent Connections: 100
• Dashboard Users: 20 concurrent

Availability Requirements:
• Uptime: 99.9% (8.76 hours downtime/year)
• Recovery Time: <5 minutes
• Data Loss: Zero data loss tolerance
"""
    
    for line in tech_spec_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())
    
    # API Specifications
    doc.add_heading('API Specifications', level=1)
    
    api_text = """
REST API Endpoints:

Health Check:
• GET /health - Returns agent health status
• Response includes component status and circuit breaker states

Agent Status:
• GET /status - Comprehensive status and metrics
• Response includes performance metrics, business metrics, and system metrics

Process Alert (A2A Protocol):
• POST /a2a/process_alert - Process security alert
• Accepts AlertMessage format
• Returns AlertProcessingResult

Agent Card (A2A Protocol):
• GET /agent-card - Returns A2A agent card
• Response includes capabilities and endpoints

HTTP Status Codes:
• 200: Success
• 400: Bad Request
• 401: Unauthorized
• 403: Forbidden
• 404: Not Found
• 500: Internal Server Error
• 503: Service Unavailable
"""
    
    for line in api_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())
    
    # Security Specifications
    doc.add_heading('Security Specifications', level=1)
    
    security_text = """
Authentication & Authorization:
• Service Account: Google Cloud service account
• API Keys: Optional API key authentication
• TLS: Required for all communications
• Permissions: Principle of least privilege

Data Protection:
• Encryption at Rest: AES-256
• Encryption in Transit: TLS 1.3
• Data Retention: Configurable (default 30 days)
• PII Handling: No PII storage

Network Security:
• Firewall: Restrictive inbound rules
• VPC: Deploy in private network
• Load Balancer: SSL termination
• DDoS Protection: Cloud-based protection

Compliance & Standards:
• ISO 27001: Information security management
• NIST Cybersecurity Framework: Security controls
• SOC 2: Security, availability, processing integrity
• GDPR: Data protection (if applicable)
"""
    
    for line in security_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())
    
    # Deployment Specifications
    doc.add_heading('Deployment Specifications', level=1)
    
    deploy_text = """
Container Specifications:
• Base Image: python:3.9-slim
• Working Directory: /app
• User: Non-root user (app)
• Ports: 8083 (agent), 8529 (dashboard)

Resource Requirements:
• Memory: 512Mi (request), 1Gi (limit)
• CPU: 250m (request), 500m (limit)
• Storage: 10Gi (models and logs)

Environment Variables:
• GOOGLE_APPLICATION_CREDENTIALS: Service account key path
• GOOGLE_CLOUD_PROJECT: Google Cloud project ID
• HYBRID_DEFENSE_CONFIG_PATH: Configuration file path
• HYBRID_DEFENSE_LOG_LEVEL: Logging level

Deployment Options:
• Development: Direct Python execution
• Production: Systemd services
• Container: Docker deployment
• Orchestration: Kubernetes deployment
"""
    
    for line in deploy_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())
    
    # Monitoring & Observability
    doc.add_heading('Monitoring & Observability', level=1)
    
    monitor_text = """
Metrics Collection:
• System Metrics: CPU, memory, disk, network
• Application Metrics: Response times, error rates, throughput
• Business Metrics: Alerts processed, actions recommended
• Custom Metrics: DQN confidence, Gemini usage

Logging:
• Log Levels: DEBUG, INFO, WARNING, ERROR, CRITICAL
• Log Format: JSON structured logging
• Log Aggregation: Centralized logging system
• Log Retention: 90 days

Alerting:
• Critical Alerts: Service down, high error rate
• Warning Alerts: Performance degradation, circuit breaker activation
• Info Alerts: Configuration changes, deployments

Health Checks:
• Agent Health: Every 30 seconds
• Component Health: DQN, Gemini, Pub/Sub
• Circuit Breaker States: Real-time monitoring
• Performance Metrics: Continuous monitoring
"""
    
    for line in monitor_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())
    
    # Testing Specifications
    doc.add_heading('Testing Specifications', level=1)
    
    testing_text = """
Unit Testing:
• Coverage: >90% code coverage
• Framework: pytest
• Mocking: External dependencies
• Performance: Benchmark critical paths

Integration Testing:
• API Testing: All endpoints
• Pub/Sub Testing: Message flow validation
• Gemini Testing: API integration
• Circuit Breaker Testing: Failure scenarios

Load Testing:
• Tools: Apache JMeter, Locust
• Scenarios: Normal load, peak load, stress test
• Metrics: Response time, throughput, error rate

Security Testing:
• Static Analysis: Code vulnerability scanning
• Dynamic Analysis: Runtime security testing
• Penetration Testing: External security assessment
"""
    
    for line in testing_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())
    
    # Future Enhancements
    doc.add_heading('Future Enhancements', level=1)
    
    future_text = """
Short-term (3-6 months):
• Multi-model Support: Additional ML models
• Enhanced Dashboard: Advanced analytics and reporting
• API Rate Limiting: Built-in rate limiting
• Configuration UI: Web-based configuration management

Medium-term (6-12 months):
• Federated Learning: Distributed model training
• Advanced Analytics: ML-based threat intelligence
• Mobile App: Mobile dashboard application
• Integration Hub: Third-party tool integrations

Long-term (12+ months):
• Autonomous Response: Automated threat containment
• Predictive Analytics: Threat prediction capabilities
• Global Deployment: Multi-region deployment
• AI Governance: Advanced AI model governance
"""
    
    for line in future_text.strip().split('\n'):
        if line.strip():
            if line.startswith('•'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())
    
    # Save document
    output_file = "deployment_package/documentation/HYBRID_DEFENSE_AGENT_TECHNICAL_SPEC.docx"
    doc.save(output_file)
    print(f"✅ DOCX document created: {output_file}")

def main():
    """Main function to create DOCX documentation"""
    
    print("📝 Creating DOCX Documentation for Hybrid Cyber Defense Agent")
    print("=" * 60)
    
    # Check if deployment package directory exists
    if not Path("deployment_package/documentation").exists():
        Path("deployment_package/documentation").mkdir(parents=True, exist_ok=True)
    
    try:
        # Create deployment guide DOCX
        create_deployment_guide_docx()
        
        # Create technical specification DOCX
        create_technical_spec_docx()
        
        print(f"\n✅ DOCX documentation created successfully!")
        print(f"📁 Location: deployment_package/documentation/")
        print(f"📄 Files created:")
        print(f"   • HYBRID_DEFENSE_AGENT_DEPLOYMENT_GUIDE.docx")
        print(f"   • HYBRID_DEFENSE_AGENT_TECHNICAL_SPEC.docx")
        
    except Exception as e:
        print(f"❌ Error creating DOCX documentation: {e}")
        return

if __name__ == "__main__":
    main()