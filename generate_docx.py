from docx import Document
from docx.shared import Pt, RGBColor
import os

def create_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('MSSP Firewall Integration Guide', 0)
    
    # Overview
    doc.add_heading('Overview', level=1)
    doc.add_paragraph(
        "This document describes the integration of external firewalls with the AI-Driven SOC Platform. "
        "The integration currently supports Palo Alto Networks (NGFW/Panorama) and Check Point (R80+ API), "
        "allowing the AI agent to manage blocking actions across hybrid fleets."
    )
    
    # Features
    doc.add_heading('Key Features', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Multi-Vendor Support: ").bold = True
    p.add_run("Unified agent interface for Palo Alto and Check Point.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Tenant-Aware: ").bold = True
    p.add_run("Credentials stored securely per-tenant with type-specific configs.")

    # Prerequisites
    doc.add_heading('Prerequisites & Requirements', level=1)
    
    doc.add_heading('1. Palo Alto Networks', level=2)
    doc.add_paragraph("Connectivity: HTTPS (443) to Management Interface.", style='List Bullet')
    doc.add_paragraph("Auth: API Key (XML API Read-Write).", style='List Bullet')
    
    doc.add_heading('2. Check Point', level=2)
    doc.add_paragraph("Connectivity: HTTPS (443) to Web API.", style='List Bullet')
    doc.add_paragraph("Auth: Username/Password (or API Key).", style='List Bullet')
    doc.add_paragraph("Context: Domain required for MDS environments.", style='List Bullet')

    # Implementation
    doc.add_heading('Implementation Details', level=1)
    
    doc.add_heading('Software Components', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Role'
    
    row = table.add_row().cells
    row[0].text = 'checkpoint_integration.py'
    row[1].text = '[NEW] Handles Check Point R80+ API (Login, Add Host, Publish).'
    
    row = table.add_row().cells
    row[0].text = 'palo_alto_integration.py'
    row[1].text = 'Handles Palo Alto XML API.'
    
    row = table.add_row().cells
    row[0].text = 'taa_a2a_mcp_agent.py'
    row[1].text = 'Unified tool logic routing based on firewall_config.type.'

    # Usage
    doc.add_heading('Configuration & Usage', level=1)
    
    doc.add_paragraph("Register a tenant with the appropriate config:")
    
    doc.add_paragraph("Check Point Example:", style='Heading 3')
    cp_sample = """{
  "firewall_config": {
    "type": "checkpoint",
    "mgmt_ip": "1.2.3.4",
    "username": "admin",
    "password": "***",
    "domain": "Global"
  }
}"""
    doc.add_paragraph(cp_sample, style='Quote')
    
    doc.add_paragraph("The AI Agent tool usage remains the same:")
    doc.add_paragraph("Tool: firewall_block_ip(tenant_id='...', ip_address='...')", style='Quote')

    # Save
    doc.save('Firewall_Integration_Guide.docx')
    print("Document saved successfully.")

if __name__ == "__main__":
    create_document()
