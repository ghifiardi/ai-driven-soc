from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_admin_guide_docx():
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Internal SOC Administrator Guide: End-to-End Onboarding')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E) # Soft Slate

    doc.add_paragraph("\nThis guide details exactly how to onboard a new customer into the AI-Driven SOC platform using our automated tools.\n")

    # Section 1: Prerequisites
    doc.add_heading('🛠️ Prerequisites', level=1)
    doc.add_paragraph("Ensure your local server is running: python3 mssp_platform_server.py", style='List Bullet')
    doc.add_paragraph("Ensure you have gcloud authenticated with the gatra-481606 project.", style='List Bullet')

    # Section 2: Onboarding Process
    doc.add_heading('🚀 The 4-Step Onboarding Process', level=1)

    doc.add_heading('Step 1: Run the Onboarding Script', level=2)
    doc.add_paragraph("Open your terminal and run the onboard_customer.py script.")
    doc.add_paragraph('python3 onboard_customer.py alpha_security "Alpha Security Corp"', style='Quote')

    doc.add_heading('Step 2: Extract Credentials', level=2)
    doc.add_paragraph("The script will output the results in your terminal. Copy the API Key generated for the customer.")
    
    doc.add_heading('Step 3: (Internal) Verify Resources', level=2)
    doc.add_paragraph("Double-check that the BigQuery datasets and Pub/Sub topics were created in your GCP Console.")

    doc.add_heading('Step 4: Handover to Customer', level=2)
    doc.add_paragraph("Provide the customer with their Tenant ID, API Key, and a copy of the AI_SOC_Customer_Integration_Guide.docx.")

    # Section 3: Management
    doc.add_heading('📋 Ongoing Management', level=1)
    
    doc.add_heading('Viewing Registered Tenants', level=2)
    doc.add_paragraph("curl http://localhost:8081/api/v1/tenants", style='Quote')

    doc.add_heading('Rotating API Keys', level=2)
    doc.add_paragraph("Manually update the config/gatra_multitenant_config.json file and restart the server.")

    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer.paragraphs[0].text = "SOC Administrator Manual | Internal Only"
    footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    filename = 'SOC_Admin_Onboarding_Manual.docx'
    doc.save(filename)
    print(f"✅ Admin Guide saved as {filename}")
    return os.path.abspath(filename)

if __name__ == "__main__":
    create_admin_guide_docx()
